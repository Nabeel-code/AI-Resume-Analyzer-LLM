# app_resume_analyzer.py
import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
import streamlit as st
from src.resume_parser import extract_text_from_resume

from src.text_utils import (
    extract_relevant_phrases,
    tfidf_similarity,
    extract_skills,
    keyword_overlap_score,
    combine_scores,
    clean_text
)
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from io import BytesIO
from docx import Document
import base64
import json

# -------------------------
# Cached model loader
# -------------------------
@st.cache_resource(show_spinner=False)
def load_embedding_model(name: str = "all-MiniLM-L6-v2"):
    model = SentenceTransformer(name)
    return model

model = load_embedding_model()

# -------------------------
# Streamlit UI
# -------------------------
st.set_page_config(page_title="Resume Analyzer — Tier-3", layout="wide")
st.title("🔎 Resume Analyzer — Tier-3 (Hybrid ATS Scoring)")

with st.sidebar:
    st.header("Controls")
    st.write("This tool uses semantic embeddings + TF-IDF + skill overlap.")
    st.markdown("---")
    st.write("Model: all-MiniLM-L6-v2 (sentence-transformers)")

uploaded = st.file_uploader("Upload Resume (.docx)", type=["docx"])
jd_text = st.text_area("Paste Job Description", height=350, placeholder="Paste full job description here (responsibilities, skills, qualifications)...")

if st.button("Run Analysis"):

    if uploaded is None:
        st.error("Upload a resume (.docx) file first.")
        st.stop()
    if jd_text.strip() == "":
        st.error("Paste a job description.")
        st.stop()

    # 1) extract resume text
    with st.spinner("Extracting resume text..."):
        resume_text = extract_text_from_resume(uploaded)
    st.subheader("Resume (preview)")
    st.write(resume_text[:2500] + ("..." if len(resume_text) > 2500 else ""))

    # 2) Cleaned full texts
    resume_clean = clean_text(resume_text)
    jd_clean = clean_text(jd_text)

    # 3) Simple TF-IDF similarity
    with st.spinner("Computing TF-IDF similarity..."):
        tfidf_sim = tfidf_similarity(resume_clean, jd_clean)

    # 4) Semantic embeddings similarity
    with st.spinner("Computing semantic similarity (embeddings)..."):
        try:
            embs = model.encode([resume_clean, jd_clean], convert_to_numpy=True)
            sem_sim = float(cosine_similarity([embs[0]], [embs[1]])[0, 0]) * 100
            sem_sim = round(sem_sim, 2)
        except Exception as e:
            st.error("Embedding model failed: " + str(e))
            sem_sim = 0.0

    # 5) Keyword/skill overlap
    with st.spinner("Extracting skills and computing overlap..."):
        resume_skills = extract_skills(resume_clean)
        jd_skills = extract_skills(jd_clean)
        keyword_sim, matched_skills = keyword_overlap_score(resume_clean, jd_clean)

    # 6) Combine into final score
    final_score = combine_scores(semantic_sim=sem_sim, tfidf_sim=tfidf_sim, keyword_sim=keyword_sim)

    # 7) Show scorecards
    st.subheader("Scores")
    col1, col2, col3 = st.columns(3)
    col1.metric("Semantic (embeddings)", f"{sem_sim:.2f}%", "")
    col2.metric("TF-IDF similarity", f"{tfidf_sim:.2f}%", "")
    col3.metric("Skill overlap", f"{keyword_sim:.2f}%", "")

    st.markdown(f"### 🔢 Combined Resume-to-JD Match: **{final_score}%**")

    # 8) Show matched and missing skills
    st.subheader("Skills found in Resume (by category)")
    st.write(resume_skills or "No lexicon matches found. Consider expanding lexicon.")

    st.subheader("Skills found in JD (by category)")
    st.write(jd_skills or "No lexicon matches found. Consider expanding lexicon.")

    st.subheader("Matched skills (intersection)")
    st.write(matched_skills or "No direct skill matches from lexicon.")

    # 9) Missing high-value skills suggestion (JD skills minus resume)
    jd_skill_words = set()
    for v in jd_skills.values():
        jd_skill_words.update(v)
    resume_skill_words = set()
    for v in resume_skills.values():
        resume_skill_words.update(v)
    missing_skills = sorted(list(jd_skill_words - resume_skill_words))

    st.subheader("Suggested skills to add")
    if missing_skills:
        st.write(", ".join(missing_skills))
    else:
        st.success("No major skill gaps detected from lexicon.")

    # 10) Auto recommendations (short)
    st.subheader("Quick Recommendations")
    recs = []
    if final_score < 50:
        recs.append("Add explicit keywords from the JD (skills & action verbs).")
    if keyword_sim < 40:
        recs.append("Mention core tooling from JD (scikit-learn, docker, aws, etc.) in skills or project bullets.")
    if sem_sim < 45:
        recs.append("Rewrite 1-2 bullets per project to focus on impact & measured outcomes.")
    if tfidf_sim < 35:
        recs.append("Add a short 'Relevant Experience' section that mirrors JD responsibilities.")
    if not recs:
        recs = ["Your resume is well-aligned — refine wording and quantify impact."]
    for r in recs:
        st.write("-", r)

    # 11) Generate DOCX report for download
    def create_docx_report(resume_name="resume", jd_title="Job Description", final_score=0.0,
                           sem_sim=0.0, tfidf_sim=0.0, keyword_sim=0.0,
                           matched_skills=None, missing_skills=None, recommendations=None):
        doc = Document()
        doc.add_heading("Resume Analyzer — Report", level=1)
        doc.add_paragraph(f"Combined Match Score: {final_score}%")
        doc.add_paragraph(f"Semantic (embeddings): {sem_sim}%")
        doc.add_paragraph(f"TF-IDF similarity: {tfidf_sim}%")
        doc.add_paragraph(f"Skill overlap: {keyword_sim}%")
        doc.add_heading("Matched Skills", level=2)
        doc.add_paragraph(", ".join(matched_skills) if matched_skills else "None")
        doc.add_heading("Missing Skills (from JD)", level=2)
        doc.add_paragraph(", ".join(missing_skills) if missing_skills else "None")
        doc.add_heading("Recommendations", level=2)
        for r in (recommendations or []):
            doc.add_paragraph("- " + r)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio

    report = create_docx_report(
        resume_name=getattr(uploaded, "name", "resume"),
        jd_title="Job",
        final_score=final_score,
        sem_sim=sem_sim,
        tfidf_sim=tfidf_sim,
        keyword_sim=keyword_sim,
        matched_skills=matched_skills,
        missing_skills=missing_skills,
        recommendations=recs
    )

    st.download_button("Download analysis report (DOCX)", data=report, file_name="resume_analysis_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
